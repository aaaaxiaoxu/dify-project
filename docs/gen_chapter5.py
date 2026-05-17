#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成第5章系统实现.docx"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

section = doc.sections[0]
section.page_width   = Cm(21)
section.page_height  = Cm(29.7)
section.left_margin  = Cm(3.18)
section.right_margin = Cm(3.18)
section.top_margin   = Cm(2.54)
section.bottom_margin= Cm(2.54)

def set_font(run, name_zh='宋体', name_en='Times New Roman', size=12, bold=False, color=None):
    run.font.name = name_en
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name_zh)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def para_spacing(p, before=0, after=6, line=300):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    pf.line_spacing = Twips(line)

def add_heading1(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, name_zh='黑体', size=16, bold=True)
    para_spacing(p, before=18, after=12, line=360)
    return p

def add_heading2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, name_zh='黑体', size=14, bold=True)
    para_spacing(p, before=12, after=6, line=320)
    return p

def add_heading3(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, name_zh='黑体', size=13, bold=True)
    para_spacing(p, before=8, after=4, line=300)
    return p

def add_body(text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    run = p.add_run(text)
    set_font(run, name_zh='宋体', size=12)
    para_spacing(p, before=0, after=6, line=300)
    return p

def add_caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, name_zh='宋体', size=10.5)
    para_spacing(p, before=2, after=10, line=240)
    return p

# ══════════════════════════════════════════
# 第5章 标题
# ══════════════════════════════════════════
add_heading1('第5章  系统实现')

add_body(
    '本章以系统各功能模块为主线，对基于Dify工作流的智能旅行日记系统的实现过程进行详细阐述。'
    '系统整体采用前后端分离架构：前端基于uni-app框架配合Vue3与uView UI组件库构建Android客户端，'
    '后端使用Python Flask轻量级Web框架提供RESTful API，持久化层选用MySQL关系型数据库，'
    '媒体文件存储依托腾讯云COS对象存储服务，AI能力通过Dify工作流平台集成大语言模型实现。'
    '各模块按照单一职责原则划分，通过统一的JWT身份认证机制串联，共同支撑旅行记录、AI智能分析、'
    '数据统计可视化、旅行报告生成与PDF导出等核心业务流程。'
)

# ══════════════════════════════════════════
# 5.1 用户管理模块
# ══════════════════════════════════════════
add_heading2('5.1  用户管理模块')

add_body(
    '用户管理模块是整个系统的安全门户，承担用户身份注册、登录鉴权与个人信息维护三大职责。'
    '在身份认证方面，系统采用JWT（JSON Web Token）实现无状态会话管理。用户提交用户名和密码后，'
    '后端首先通过Flask-JWT-Extended扩展对请求体进行参数校验，随后在MySQL的users表中查询匹配记录，'
    '使用bcrypt算法对提交密码与数据库中存储的哈希值进行比对验证。bcrypt算法内置加盐机制，'
    '相同明文每次生成的哈希值均不相同，有效抵御彩虹表攻击，确保用户密码在数据库泄露情形下仍具备较强的安全性。'
)
add_body(
    '身份验证通过后，后端调用create_access_token接口签发JWT令牌，令牌载荷中包含用户ID、角色信息及'
    '过期时间戳等声明字段。前端将令牌存储于uni-app的本地持久化存储中，后续每次发起API请求时，'
    '在HTTP请求头的Authorization字段中以"Bearer <token>"格式携带令牌。后端通过@jwt_required()装饰器'
    '拦截受保护路由，自动完成令牌有效性校验与身份解析，无需在每个接口中重复实现鉴权逻辑。'
    '若令牌过期或被篡改，后端统一返回401状态码，前端捕获后跳转至登录页面引导用户重新认证。'
)
add_body(
    '在个人信息管理方面，用户可在个人中心页面查看并修改昵称、手机号等基础资料。头像上传功能由前端'
    '调用uni.chooseImage接口唤起系统相册或相机，获取图片临时路径后，通过uni.uploadFile发起multipart/form-data'
    '格式的文件上传请求至后端/api/file/upload接口。后端接收文件后利用腾讯云COS Python SDK将文件流'
    '写入预设的Bucket路径，COS返回的持久访问URL被写回users表的avatar_url字段，前端下次加载时即可'
    '通过该URL直接渲染头像图片，原始文件不占用服务器本地磁盘空间。'
)
add_body(
    '此外，系统设计了密码重置流程：用户在遗忘密码时可通过手机号验证码完成身份核实，'
    '验证通过后允许设置新密码，新密码同样经bcrypt哈希处理后写入数据库，整个流程无需暴露旧密码。'
    '登录界面如图5.1所示，页面设计简洁，突出品牌标识与输入区域，并在底部提供注册和忘记密码的快速入口。'
)
add_caption('图5.1  用户登录界面')

# ══════════════════════════════════════════
# 5.2 日记管理模块
# ══════════════════════════════════════════
add_heading2('5.2  日记管理模块')

add_body(
    '日记管理模块是系统的核心数据入口，为用户提供旅行记录的全生命周期管理能力，涵盖日记的创建、'
    '编辑、草稿保存、删除及浏览等完整操作链路。该模块在数据模型设计上采用主从表结构：主表diaries'
    '存储日记的基础元数据，包括标题（title）、富文本内容（content）、旅行地点文字（location）、'
    '地理坐标（latitude/longitude）、情绪标签（mood）、旅行日期（date）、是否草稿（is_draft）等字段；'
    '关联表diary_images和diary_videos分别存储日记附带的图片和视频资源URL，与主表通过diary_id进行外键关联，'
    '支持一篇日记对应多张图片和多段视频的灵活媒体组合。'
)
add_body(
    '在内容编辑方面，前端集成了功能完整的富文本编辑器，用户可以对日记正文进行加粗、斜体、标题层级、'
    '有序/无序列表、图片嵌入等格式化操作，满足不同写作习惯的表达需求。图片上传采用"先上传后引用"策略：'
    '用户在编辑器内插入图片时，前端立即将图片文件通过multipart请求上传至/api/file/upload接口，'
    '后端经腾讯云COS SDK存储后返回持久URL，编辑器随即将该URL插入到富文本HTML的img标签src属性中。'
    '最终保存日记时，content字段存储的是包含完整COS URL的HTML字符串，不依赖任何本地临时路径，'
    '确保内容在多端访问时的一致性与持久性。视频资源同理，通过video标签嵌入富文本内容中。'
)
add_body(
    '草稿自动保存机制通过前端定时器实现：用户在编辑过程中，系统每隔一定时间自动将当前内容以is_draft=True'
    '的状态提交至后端保存接口，即使应用意外退出，再次进入编辑页时也能自动恢复上次的草稿内容，'
    '有效防止因意外情况造成的内容丢失。用户主动点击发布时，后端将is_draft字段更新为False，'
    '该日记随即出现在日记列表和地图轨迹等所有需要"已发布"状态的功能模块中。'
)
add_body(
    '地理位置信息的获取在编辑页进入时自动触发：前端调用uni.getLocation接口请求用户当前的GPS坐标，'
    '获取成功后将经纬度传至后端的腾讯地图逆地理编码接口，解析出省、市、区三级行政区文字，'
    '拼装为可读地名后同步填入location输入框供用户确认或修改，原始经纬度则始终保留在latitude/longitude'
    '字段，为地图轨迹模块的标点渲染提供精确的坐标依据。日记编辑界面如图5.2所示，日记列表界面如图5.3所示。'
)
add_caption('图5.2  日记编辑界面')
add_caption('图5.3  日记列表界面')

# ══════════════════════════════════════════
# 5.3 AI智能分析模块（重点）
# ══════════════════════════════════════════
add_heading2('5.3  AI智能分析模块')

add_body(
    '本系统将AI智能分析定位为差异化核心功能，旨在突破传统旅行日记"仅作记录"的局限，'
    '为每篇日记提供情感解读、关键词提炼、旅行偏好识别、写作风格评估和个性化旅行建议等多维度智能增值服务。'
    '在技术选型上，系统引入Dify大模型工作流平台作为主力AI引擎，同时在后端实现了一套基于jieba分词'
    '的本地分析算法作为降级保障，形成"大模型优先、本地算法兜底"的双通道冗余架构，在保证分析质量的'
    '同时兼顾系统在网络异常或API配额耗尽情形下的基础可用性。'
)

add_heading3('5.3.1  Dify工作流集成与多模态分析')

add_body(
    'Dify是一个开源的大语言模型应用开发平台，其核心优势在于通过可视化的节点编排界面，'
    '将Prompt模板设计、变量注入、多模态输入处理、条件分支、大模型调用和结构化输出解析等'
    '复杂流程串联为可复用的工作流，并以标准HTTP API对外暴露，调用方无需感知内部实现细节。'
    '本系统在Dify平台上共部署了三个独立工作流，分别承担不同的AI任务：日记智能分析工作流、'
    '写作建议工作流和旅行总结报告生成工作流。三个工作流均通过DifyClient类统一管理，'
    '各自维护独立的API密钥和服务地址配置，互不干扰。'
)
add_body(
    '日记智能分析工作流是AI分析模块的核心。当用户在日记详情页点击"AI分析"按钮时，'
    '前端向后端发起POST /api/ai/analysis请求，携带diary_id参数。后端接口首先查询日记内容，'
    '调用html_to_plain_text工具函数将富文本HTML剥离标签转为纯文本，以便大模型正确理解语义；'
    '同时从diary_images和diary_videos关联表中收集所有已上传的媒体资源URL，'
    '构造成Dify multimodal输入所要求的标准结构：图片以{"type": "image", "transfer_method": "remote_url", "url": "..."}格式组装，'
    '视频以{"type": "video", "transfer_method": "remote_url", "url": "..."}格式组装，'
    '与日记文本一并作为工作流的inputs参数，通过HTTP POST请求发送至Dify API端点/workflows/run。'
    'Dify工作流节点在收到请求后，以remote_url方式直接拉取COS上的图片和视频内容，'
    '结合文本进行跨模态联合理解，能够识别照片中的景点、活动场景，以及视频中的环境氛围，'
    '使分析结果更贴合日记的真实情境，而非仅凭文字推断。'
)
add_caption('图5.4  AI智能分析处理流程')

add_body(
    '请求采用blocking（同步阻塞）响应模式，后端等待Dify工作流执行完毕后一次性获取完整结果。'
    'response_mode设置为blocking的优势在于实现简单、结果确定，适合本系统对分析结果完整性要求'
    '较高的场景，超时时间设置为60秒以应对大模型推理延迟。DifyClient的_call_workflow方法'
    '封装了完整的HTTP请求构建、异常捕获和错误日志记录逻辑，如图5.5所示。'
)
add_caption('图5.5  Dify工作流调用核心代码')

add_heading3('5.3.2  分析结果结构与持久化')

add_body(
    'Dify工作流的输出节点配置为返回一个名为analysis_result的字段，其内容是由大模型按照'
    '预设JSON Schema生成的结构化JSON字符串。由于大模型输出存在一定随机性，'
    '返回的JSON有时会被包裹在Markdown代码块（```json ... ```）中，'
    '系统通过_parse_json_output方法进行统一的清洗处理：使用正则表达式剥离代码块标记，'
    '再调用json.loads进行反序列化，对解析失败的情况降级为返回原始字符串以保留分析内容。'
)
add_body(
    '解析成功后，系统从JSON对象中提取以下七个标准化字段：emotion_label（情感标签，'
    '如"积极""平和""略感失落"等描述性词语）、emotion_score（情绪评分，浮点数，'
    '由大模型根据文本情绪强度赋值，系统对其归一化处理限定在-1.0至1.0区间）、'
    'emotion_analysis（情感分析详细描述，解释情绪判断的依据）、keywords（关键词数组，'
    '提炼日记核心主题词）、travel_advice（个性化旅行建议，基于日记内容给出下次旅行方向）、'
    'memory_point（记忆点，提炼本次旅行中最值得铭记的时刻或细节）、'
    'writing_suggestion（写作建议，从叙事角度给出日记内容的改进提示）。'
    '上述字段的解析与映射逻辑如图5.6所示。'
)
add_body(
    '分析完成后，结果被持久化至ai_analysis表。系统先查询是否存在该日记的历史分析记录：'
    '若存在则执行UPDATE更新各字段，若不存在则INSERT新记录，并通过IntegrityError异常'
    '捕获处理并发写入冲突。ai_analysis表通过diary_id字段与diaries表建立一对一关联关系，'
    '设置UNIQUE约束防止重复写入。后续数据统计、报告生成等模块在需要分析数据时直接从该表'
    '读取缓存结果，避免对相同日记重复调用AI接口，显著节约大模型API调用成本。'
)
add_caption('图5.6  Dify返回结果解析与字段映射代码')

add_heading3('5.3.3  本地算法降级策略')

add_body(
    '为保障在Dify服务不可用（网络超时、API密钥未配置或调用配额耗尽）时系统仍能提供基础'
    '分析能力，后端在routes/ai.py中实现了一套完整的本地分析算法作为降级方案。'
    '该方案基于jieba中文分词库构建，无需调用任何外部API，运行时延极低，'
    '能在毫秒级完成对日记内容的初步分析。'
)
add_body(
    '本地算法由六个功能函数组成，协同完成综合分析任务。首先是情感分析函数analyze_emotion：'
    '系统预先维护三类情感词典，正面情感词典包含"开心""愉快""兴奋""幸福"等15个高频积极词汇，'
    '负面情感词典包含"难过""沮丧""焦虑""郁闷"等13个消极词汇，中性词典包含"平静""安宁"'
    '"悠闲"等8个平和词汇（词典定义如图5.7所示）。jieba对日记文本完成分词后，函数统计各类'
    '情感词的出现频次，计算正面词占比与负面词占比之差作为情绪评分的基准值，'
    '归一化至[-1.0, 1.0]区间后输出，同时根据各类情感词的相对比例确定情感标签。'
)
add_caption('图5.7  本地分析知识库——情感与地点类型词典')

add_body(
    '关键词提取函数extract_keywords_tfidf基于jieba内置的TF-IDF算法实现。'
    'TF-IDF综合考量词语在当前文档中的出现频率（TF）与在全局语料中的稀缺程度（IDF），'
    '能有效过滤"的""了""和"等高频虚词，优先提取具有实际语义价值的名词和动词，'
    '默认返回权重最高的前5个关键词，为日记提供直观的主题摘要。'
    '地点类型识别函数identify_location_type通过预定义的五类地点词典（自然风光、城市景观、'
    '历史文化、现代建筑、乡村田园）对分词结果进行词频匹配，识别日记主要描写的场景类型。'
    '活动类型识别函数identify_activities同理，通过五类活动词典（观光游览、美食体验、'
    '休闲娱乐、户外运动、文化体验）识别旅行中的主要活动。最终generate_travel_advice函数'
    '根据识别出的地点类型和活动类型，从预设的建议模板库中随机抽取一条地点建议，'
    '并拼接对应的活动建议，形成个性化的旅行推荐文字输出给用户。'
)
add_body(
    '写作建议方面，系统在Dify分析通道和本地分析通道之外，还设计了独立的写作建议增强机制：'
    '_merge_writing_feedback函数在获取主分析结果后，额外调用Dify的写作建议工作流（仅接收'
    '纯文本输入），若该工作流返回有效建议则覆盖本地生成的写作建议；若Dify写作工作流同样'
    '不可用，则本地analyze_writing_style函数通过统计文本字符数和平均句长，'
    '从"简洁""适中""详细"三个维度对写作风格作出初步判断并给出改进方向。'
    '这一设计使写作建议功能形成了三级保障链路，最大程度保证用户始终能收到有参考价值的'
    '写作反馈。AI分析结果的前端展示效果如图5.8所示。'
)
add_caption('图5.8  AI智能分析结果界面')

# ══════════════════════════════════════════
# 5.4 数据分析与可视化模块（重点）
# ══════════════════════════════════════════
add_heading2('5.4  数据分析与可视化模块')

add_body(
    '数据分析与可视化模块在ai_analysis表积累的情绪评分数据基础上，为用户提供宏观维度的'
    '情感洞察与旅行行为统计，帮助用户从更高视角审视自己的旅行情感轨迹与记录习惯。'
    '该模块的后端实现集中在routes/stats.py文件中，提供两个核心统计接口；'
    '前端统计分析页面（stats.vue）基于uni-app的Canvas API手动绘制饼图与双折线图，'
    '实现无第三方图表库依赖的轻量级可视化方案，减少APP安装包体积。'
)

add_heading3('5.4.1  情感分布统计')

add_body(
    '情感分布接口GET /api/stats/emotion-distribution以当前登录用户的身份标识为过滤条件，'
    '通过SQLAlchemy ORM联表查询diaries表和ai_analysis表，筛选出所有is_draft=False'
    '（已正式发布）的日记所对应的情绪评分记录。查询结果按照预定义的五档评分区间进行归类：'
    '强烈积极（emotion_score≥0.6）、偏积极（0.2≤emotion_score<0.6）、'
    '中性（-0.2≤emotion_score<0.2）、偏消极（-0.6≤emotion_score<-0.2）、'
    '强烈消极（emotion_score<-0.6）。对于尚未完成AI分析、情绪评分为NULL的日记，'
    '单独归入"待分析"类别，在饼图中以灰色显示，提示用户仍有日记未进行AI分析处理。'
)
add_body(
    '接口在完成分组统计后，同步计算所有有效评分（非NULL值）的算术平均值作为整体情绪基准分，'
    '与各区间的计数、占比数据一并返回给前端。响应体中的items数组按固定顺序排列六个分组，'
    '每个分组包含emotion（分组名称）、count（日记数量）、range（评分区间描述）三个字段，'
    '前端据此绘制饼图并生成图例说明。区间划分与聚合逻辑的核心实现如图5.9所示。'
)
add_caption('图5.9  情感分布统计接口核心代码')

add_body(
    '前端stats.vue页面顶部的概览卡片区域直观展示三项摘要指标：日记总数（取items中所有分组'
    '数量之和）、主要情绪区间（取count最大的分组名称）、平均情绪评分（将浮点数格式化为'
    '带正负号的两位小数字符串）。饼图区域使用uni-app的canvas组件，前端JavaScript根据'
    '各分组数量计算对应扇形的起止弧度，调用ctx.arc()和ctx.fill()方法逐段绘制，'
    '不同情绪区间分配差异化配色（强烈积极为深蓝、偏积极为天蓝、中性为灰色等）以便区分。'
    '饼图下方的图例列表逐行显示各区间的情绪名称、评分范围和日记篇数占比，'
    '辅助用户理解图形含义。情感分布统计界面效果如图5.10所示。'
)
add_caption('图5.10  情感分布统计界面')

add_heading3('5.4.2  情绪波动趋势')

add_body(
    '情绪波动趋势接口GET /api/stats/emotion-trend通过查询参数period控制时间范围，'
    '支持"7d"（近7天）、"30d"（近30天）和"all"（全部历史）三种模式。'
    '在时间粒度设计上，系统针对不同数据量规模自适应切换聚合单位：当period为7d或30d时，'
    '统计粒度固定为"按天"，每个日历日作为一个数据点；当period为all且日记历史跨度超过90天时，'
    '系统自动将粒度升级为"按月"聚合，避免折线图出现数百个密集数据点导致视觉杂乱、'
    '难以识别整体趋势的问题。粒度切换逻辑通过对比最早与最晚日记日期之差实现，'
    '完全动态，无需用户手动选择。'
)
add_body(
    '接口查询逻辑采用时间桶（bucket）预填充策略：后端首先根据时间范围和粒度在内存中'
    '构建一个有序字典，以日期字符串（按天为"YYYY-MM-DD"，按月为"YYYY-MM"）为键，'
    '预填充count=0、score_sum=0、score_cnt=0的初始状态。随后逐行遍历数据库查询结果，'
    '将每条日记的情绪评分累加至对应时间桶，最终按桶聚合计算平均情绪评分。'
    '此策略的优势在于：即使某个时间段内没有日记记录，该时间点仍会出现在返回的dates数组中'
    '（对应count=0、score=null），前端折线图不会出现时间轴断裂，能够完整呈现连续的时间序列。'
    '情绪趋势统计接口的粒度自适应实现如图5.11所示。'
)
add_caption('图5.11  情绪趋势统计接口核心代码')

add_body(
    '前端折线图采用双Y轴设计以同时展示两个量纲不同的指标：左侧Y轴对应日记数量（整数，'
    '范围0至最大日记数+1），右侧Y轴对应情绪评分（浮点数，固定范围-1.2至1.2）。'
    '蓝色折线描绘各时间点的日记数量变化，红色折线描绘情绪评分波动，'
    '在情绪评分Y轴的0刻度处绘制一条灰色虚线基准线，用户可据此直观判断某段时间内'
    '情绪整体处于正值区（偏积极）还是负值区（偏消极）。'
    '对于scores数组中的null值（该时间点无分析记录），前端在绘制折线时跳过该点，'
    '避免将0值误导为"中性情绪"。'
    '页面顶部的周期切换标签支持用户在三个时间范围间自由切换，切换时重新请求接口并刷新画布。'
    '情绪波动趋势折线图的实际展示效果如图5.12所示。'
)
add_caption('图5.12  情绪波动趋势统计界面')

# ══════════════════════════════════════════
# 5.5 智能旅行总结与导出模块（重点）
# ══════════════════════════════════════════
add_heading2('5.5  智能旅行总结与导出模块')

add_body(
    '智能旅行总结与导出模块是本系统AI能力的集中呈现，也是区别于普通日记应用的核心竞争力之一。'
    '该模块将用户在某段旅行时间内的所有日记数据、AI分析结果、地理位置记录和多媒体资源进行深度聚合，'
    '经由大模型进行叙事化生成处理，最终输出一份结构完整、内容丰富、具有个人风格的旅行总结报告，'
    '并支持一键导出为格式规范的PDF文件，便于长期保存和社交分享。'
    '该功能的后端实现分布于routes/report.py和utils/report_pdf.py两个文件中，'
    '前端由report.vue页面提供交互界面。'
)

add_heading3('5.5.1  报告数据聚合')

add_body(
    '用户在报告页选择时间范围后，前端向POST /api/report/generate接口提交包含preset'
    '（预设时间范围）或start_date/end_date（自定义日期）参数的请求体。'
    '后端_resolve_report_range函数负责参数解析：预设模式下，"7d"换算为今日往前推6天，'
    '"30d"换算为往前推29天，"all"则通过MIN(date)和MAX(date)聚合查询动态确定边界；'
    '自定义模式下对日期格式进行ISO 8601格式校验，并检查start_date不晚于end_date的逻辑约束。'
    '确定日期区间后，通过SQLAlchemy执行LEFT OUTER JOIN查询，'
    '将diaries表与ai_analysis表连接，按日期升序返回区间内所有已发布日记及其分析记录。'
)
add_body(
    '_build_report_context函数接收查询结果后，对每篇日记调用_load_analysis_snapshot'
    '方法获取分析快照。若某篇日记已有ai_analysis记录且字段完整，则直接读取；'
    '若分析记录缺失或关键字段（emotion_label、memory_point、keywords）均为空，'
    '则临时调用本地analyze_diary_content函数生成分析快照，确保每篇日记都有可用的分析数据'
    '参与报告聚合，避免因部分日记未进行AI分析而导致报告内容缺失。'
)
add_body(
    '聚合过程涵盖多个维度的统计与筛选：使用Counter对地点字符串和关键词分别进行词频统计，'
    '取访问次数最多的Top 8地点和Top 10关键词；按情绪分组统计积极、中性、消极三类日记的'
    '数量占比；通过_select_highlights函数按"记忆点完整度→情绪评分绝对值→关键词数量→日期"'
    '四维优先级排序，筛选出最具代表性的Top 6篇亮点日记；通过geo_utils模块的'
    'route_distance_km_from_diaries函数基于Haversine公式估算行程总距离；'
    '通过_select_report_images函数从每篇日记的图片关联表中各取第一张照片，'
    '汇集为报告封面图片集。整体聚合逻辑如图5.13时序图所示，核心代码如图5.14所示。'
)
add_caption('图5.13  智能旅行报告生成时序')
add_caption('图5.14  报告上下文聚合核心代码')

add_heading3('5.5.2  Dify报告工作流与本地兜底')

add_body(
    '完成数据聚合后，后端将报告上下文对象（report_context）通过json.dumps序列化为JSON字符串，'
    '连同start_date、end_date和report_style（报告风格，默认为"warm"温暖风格，可选"professional"专业风格）'
    '作为inputs参数，调用DifyClient.generate_travel_report方法将数据提交给Dify旅行总结'
    '生成工作流。该工作流内部通过精心设计的Prompt模板指导大模型将结构化数据转化为自然语言叙述，'
    '生成包含以下字段的完整报告对象：report_title（报告标题，结合时间范围和主要旅行地点生成）、'
    'report_subtitle（副标题，概括日记数量和整体情绪基调）、summary（旅行摘要，'
    '介绍旅程时长、地点数量、里程和情绪均值）、highlights（亮点列表，逐条描述Top高光日记）、'
    'emotion_review（情感回顾，分析整体情绪分布和情绪变化的主要原因）、'
    'travel_preferences（旅行偏好洞察，基于地点、关键词和情绪数据总结用户的旅行风格）、'
    'next_trip_suggestions（下次旅行建议，结合历史偏好给出具体可行的行程方向）、'
    'memory_quote（记忆金句，从亮点日记中提炼最有感召力的一段话或记忆点文字）。'
)
add_body(
    '若Dify报告工作流调用失败或未配置，系统自动回退至_build_local_report函数。'
    '该函数不依赖大模型，完全通过预定义的规则和模板拼装生成与Dify输出同结构的报告对象：'
    '标题由时间范围文字和Top地点名称拼接而成，摘要通过字符串模板填入统计数值，'
    '亮点列表由_build_highlight_lines函数将每条亮点日记格式化为"日期+地点+标题+情绪分+记忆点"'
    '的固定叙述句式，情感回顾由_build_emotion_review函数根据积极/消极日记数量占比和平均分'
    '选择对应的描述模板组合生成，旅行偏好由_build_travel_preferences函数依据地点频次、'
    '关键词和情绪均值逐条拼装分析结论，下次旅行建议由_build_next_trip_suggestions函数'
    '优先提取各日记的travel_advice字段内容，不足时补充默认建议句式。'
    '_merge_report_payload函数负责将Dify结果与本地结果合并，以Dify内容为优先，'
    '对Dify未返回或为空的字段以本地生成内容兜底填充，保证最终报告对象的字段完整性。'
    '本地兜底报告生成的核心代码如图5.15所示，前端报告展示界面如图5.16所示。'
)
add_caption('图5.15  本地兜底报告生成核心代码')
add_caption('图5.16  智能旅行总结报告界面')

add_heading3('5.5.3  PDF导出实现')

add_body(
    '旅行总结报告的PDF导出功能由后端utils/report_pdf.py模块实现，核心依赖Python的reportlab库。'
    'reportlab提供底层的PDF文档对象模型，支持以编程方式精确控制页面布局、字体样式、'
    '图片嵌入和段落排版，非常适合生成结构固定但内容动态的文档。'
)
add_body(
    '用户在前端点击"导出PDF"按钮后，前端将当前页面展示的完整报告数据（包括period时间区间、'
    'summary_stats统计数据、report各文本字段、report_images图片URL列表）序列化为JSON，'
    '提交至POST /api/report/export-pdf接口。后端_normalize_export_bundle函数'
    '首先对入参进行严格的字段校验和类型归一化处理，确保必填字段（报告标题、时间范围、摘要）'
    '均存在且非空，对类型不符的字段以默认值替换，防止reportlab在生成过程中因空值或'
    '类型错误引发运行时异常。'
)
add_body(
    'build_travel_report_pdf函数接收归一化后的数据，构建PDF文档结构：封面页以大号字体居中'
    '排版报告标题和副标题，下方插入时间范围和摘要统计数据；图片集页逐张从腾讯云COS下载'
    '旅行照片，使用reportlab的Image对象按网格布局嵌入页面，每张照片下方附上日记日期和地点'
    '文字说明；正文各章节分别排版summary摘要、highlights亮点列表、emotion_review情感回顾、'
    'travel_preferences旅行偏好、next_trip_suggestions下次建议，各章节间以横线分隔；'
    '末页以引号格式大字体展示memory_quote记忆金句，形成完整的视觉收尾。'
    '中文排版方面，系统通过@lru_cache缓存字体加载结果，使用内置的中文字体资源注册'
    'reportlab字体，避免PDF中出现中文字符乱码或缺字问题。'
)
add_body(
    'PDF文件生成完毕后保存至服务器uploads/report_exports目录，文件名由'
    '"travel_report_"前缀、用户ID、时间戳和UUID随机后缀四部分拼接而成，'
    '格式为travel_report_{user_id}_{timestamp}_{uuid8}.pdf，确保每次导出生成唯一文件名。'
    '下载接口GET /api/report/download/<file_name>在响应前执行双重安全校验：'
    '一是检查文件名后缀必须为".pdf"；二是验证文件名前缀必须以"travel_report_{当前用户ID}_"'
    '开头，确保用户只能下载属于自己账号生成的报告文件，从接口设计层面防止用户通过猜测'
    '文件名越权下载他人报告。校验通过后调用Flask的send_file函数以attachment模式返回文件，'
    '触发浏览器或客户端的文件下载行为。相关接口实现如图5.17所示，PDF导出效果如图5.18所示。'
)
add_caption('图5.17  PDF导出接口核心代码')
add_caption('图5.18  旅行报告PDF导出效果')

# ══════════════════════════════════════════
# 5.6 地图轨迹与位置服务模块
# ══════════════════════════════════════════
add_heading2('5.6  地图轨迹与位置服务模块')

add_body(
    '地图轨迹与位置服务模块为用户提供直观的旅行足迹可视化功能，将抽象的地理坐标数据'
    '转化为地图上可交互的轨迹标记，帮助用户以空间视角回顾自己走过的旅行路径。'
    '前端地图页面（map.vue）通过uni-app集成腾讯地图SDK，'
    '在应用启动时调用后端GET /api/map/locations接口，获取当前用户所有已发布日记中'
    '记录的经纬度数据集。后端接口查询diaries表中latitude和longitude字段非空的记录，'
    '按日期降序返回包含日记ID、标题、日期、情绪标签、经度和纬度的精简数据列表，'
    '前端将每条记录渲染为腾讯地图上的一个标记点（marker）。'
)
add_body(
    '标记点的样式根据日记的情绪标签动态调整颜色，积极情绪显示为暖色系图标，'
    '消极情绪显示为冷色系图标，未分析的日记使用默认灰色图标，'
    '使用户在地图视角下也能感知不同旅行时期的情绪状态。'
    '点击某个标记点后，地图弹出信息气泡窗口，展示该篇日记的标题、旅行日期和情绪标签摘要；'
    '用户可进一步点击气泡中的"查看详情"按钮，跳转至对应日记的详情页面阅读完整内容。'
)
add_body(
    '在位置数据的采集环节，日记编辑页面进入时自动触发位置获取流程：'
    '调用uni.getLocation接口请求用户授权后获取当前GPS经纬度坐标，'
    '随即将坐标传至后端的腾讯地图逆地理编码接口（/api/geocode/reverse），'
    '后端转发至腾讯地图Web Service API，解析返回的省市区行政区划文字，'
    '拼装为"城市+区县"格式的简短地名字符串，自动填入日记编辑表单的地点输入框中，'
    '用户可直接使用或手动修改为更具描述性的地点名称。原始经纬度坐标与文本地名'
    '各自独立存储，两者共同支撑地图标点和文本展示两种不同的使用场景，互不覆盖。'
    '旅行足迹地图的实际效果如图5.19所示。'
)
add_caption('图5.19  旅行足迹地图界面')

# ══════════════════════════════════════════
# 5.7 社交分享模块
# ══════════════════════════════════════════
add_heading2('5.7  社交分享模块')

add_body(
    '社交分享模块允许用户将精彩的旅行日记分享给朋友和家人，在保护隐私的前提下拓展旅行记录的'
    '传播范围。分享功能的实现采用令牌（Token）机制：用户在日记详情页点击分享按钮后，'
    '后端为该篇日记生成一个全局唯一的随机令牌字符串（使用uuid4生成，长度32位），'
    '与日记ID、分享创建者用户ID、过期时间、最大访问次数等字段一同写入share_links表。'
    '系统返回形如/share/{token}的分享链接，用户可将链接复制至微信、微博等社交平台传播。'
)
add_body(
    '访问分享链接无需登录认证：访客通过链接进入后，后端从share_links表查询对应令牌的记录，'
    '依次校验以下条件：令牌是否存在、是否已过期（通过对比当前时间与expires_at字段）、'
    '实际访问次数是否已达到max_visits上限。三项校验全部通过后，接口返回日记的公开内容，'
    '包括标题、正文、图片和情绪标签，同时将该令牌的访问计数加一写回数据库。'
    '若任一校验失败，接口返回相应的错误提示，引导访客联系分享者更新链接。'
)
add_body(
    '在权限控制方面，用户可在分享管理页面（manage.vue）查看自己创建的所有分享记录，'
    '包括各链接的当前访问次数、剩余有效次数和过期时间，并可随时手动撤销某个分享，'
    '撤销后对应令牌立即失效，持有旧链接的访客将无法再访问。'
    '这一设计在提供便捷分享功能的同时，将分享权限的控制权完全交还给日记作者，'
    '有效保护用户的内容隐私安全。分享管理界面如图5.20所示。'
)
add_caption('图5.20  社交分享管理界面')

# ══════════════════════════════════════════
# 5.8 系统管理模块
# ══════════════════════════════════════════
add_heading2('5.8  系统管理模块')

add_body(
    '系统管理模块为管理员提供后台运营管理工具，承担用户账号管理和日记内容审核两大职责，'
    '确保平台内容的合规性与用户账号的安全性。在权限设计上，系统通过在users表中设置'
    'role字段区分普通用户（role="user"）与管理员（role="admin"），'
    '后端在routes/admin.py中使用admin_required自定义装饰器，对所有管理接口进行二次鉴权：'
    '在JWT鉴权通过的基础上，进一步从令牌载荷中读取当前用户的role字段，'
    '确认其值为"admin"后方允许执行操作，否则返回403 Forbidden响应，'
    '从接口层面杜绝普通用户通过篡改请求直接访问管理功能的越权风险。'
)
add_body(
    '用户管理功能允许管理员以分页方式浏览全体注册用户列表，列表字段包含用户ID、用户名、'
    '昵称、注册时间、账号状态（正常/禁用）及角色等基础信息。管理员可对违规用户执行'
    '禁用操作，后端将users表中对应记录的is_active字段置为False，'
    '该用户后续尝试登录时系统在密码校验通过后还会额外检查is_active状态，'
    '状态为False时拒绝签发JWT令牌并返回"账号已被禁用"的提示，'
    '实现对违规用户访问权限的即时切断，无需删除账号数据。'
)
add_body(
    '日记内容管理功能允许管理员查看所有用户的公开日记列表，支持按用户ID、日记标题关键字'
    '进行过滤检索。发现违规内容时，管理员可执行删除操作，后端级联删除日记主记录及其关联的'
    'diary_images、diary_videos、ai_analysis、share_links等关联表数据，'
    '确保数据库中不残留孤立记录。此外，前端管理页面（admin.vue）还展示平台级别的统计概况，'
    '包括总用户数、总日记数、今日新增用户数和今日新增日记数等运营指标，'
    '帮助管理员及时掌握平台使用动态。系统管理后台界面如图5.21所示。'
)
add_caption('图5.21  系统管理后台界面')

# ══════════════════════════════════════════
# 附：配图清单
# ══════════════════════════════════════════
doc.add_page_break()
add_heading2('附：配图清单（定稿时可删除本页）')

headers = ['图号', '类型', '内容说明', '文件来源']
col_widths = [Cm(2), Cm(2.5), Cm(9), Cm(3.5)]

table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
for i, w in enumerate(col_widths):
    for cell in table.columns[i].cells:
        cell.width = w

hrow = table.rows[0]
for i, h in enumerate(headers):
    cell = hrow.cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(h)
    set_font(run, name_zh='黑体', size=10.5, bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'D9E1F2')
    tcPr.append(shd)

rows_data = [
    ('图5.1',  '页面截图', '用户登录页',                                           '手机截图'),
    ('图5.2',  '页面截图', '日记编辑页（富文本+情绪标签）',                        '手机截图'),
    ('图5.3',  '页面截图', '日记列表页',                                            '手机截图'),
    ('图5.4',  '流程图',   'AI分析主流程（Dify工作流→本地降级完整链路）',           '自绘（draw.io）'),
    ('图5.5',  '代码截图', 'dify_client.py 第35～105行（_run_workflow + _call_workflow）', 'VSCode截图'),
    ('图5.6',  '代码截图', 'dify_client.py 第196～245行（JSON解析与字段映射）',     'VSCode截图'),
    ('图5.7',  '代码截图', 'ai.py 第15～38行（情感词典与地点类型词典定义）',        'VSCode截图'),
    ('图5.8',  '页面截图', 'AI分析结果页（情感标签/关键词/旅行建议/写作建议）',     '手机截图'),
    ('图5.9',  '代码截图', 'stats.py 第59～113行（emotion_distribution接口）',     'VSCode截图'),
    ('图5.10', '页面截图', '情感分布饼图页（附图例，各区间占比）',                  '手机截图'),
    ('图5.11', '代码截图', 'stats.py 第116～225行（emotion_trend粒度自适应逻辑）', 'VSCode截图'),
    ('图5.12', '页面截图', '情绪波动趋势折线图（双折线+基准线）',                   '手机截图'),
    ('图5.13', '时序图',   '报告生成时序（前端→聚合→Dify报告工作流→本地兜底→返回）','自绘（draw.io）'),
    ('图5.14', '代码截图', 'report.py 第294～376行（_build_report_context函数）',  'VSCode截图'),
    ('图5.15', '代码截图', 'report.py 第496～534行（_build_local_report函数）',    'VSCode截图'),
    ('图5.16', '页面截图', '智能旅行总结报告页（图片+统计+亮点+情感回顾）',         '手机截图'),
    ('图5.17', '代码截图', 'report.py 第709～765行（export-pdf + download接口）',  'VSCode截图'),
    ('图5.18', 'PDF截图',  '导出PDF成品预览（完整排版效果）',                       'PDF阅读器截图'),
    ('图5.19', '页面截图', '旅行足迹地图（多点标记+情绪色彩）',                     '手机截图'),
    ('图5.20', '页面截图', '社交分享管理页（链接列表+访问统计）',                   '手机截图'),
    ('图5.21', '页面截图', '系统管理后台（用户管理+日记管理+平台统计）',            '手机截图'),
]

for r in rows_data:
    row = table.add_row()
    for i, text in enumerate(r):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(text)
        set_font(run, name_zh='宋体', size=10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

out_path = '/Users/xuliwei/github-project/dify-project/docs/第5章_系统实现.docx'
doc.save(out_path)
print(f'已生成：{out_path}')
